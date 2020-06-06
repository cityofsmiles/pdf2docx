#!/usr/bin/env python

# go 'cd /root/storage/emulated/0/GNURoot/home/Scripts/fedora/pythondocx && python3 pdf2docx.py && termux-open demo.docx'


import os
from docx import Document
from docx.shared import Inches
import sys
import glob
import shutil
from wand.image import Image
from PyPDF2 import PdfFileReader
import argparse

status=0

def parseArguments():
    # Create argument parser
    parser = argparse.ArgumentParser()

    # Positional mandatory arguments
    parser.add_argument("inputPdf", help="Filename of the input PDF.", type=str)

    # Optional arguments
    parser.add_argument("-s", "--papersize", help="Papersize. Valid values are a4, letter, and legal.", type=str, default='')
    parser.add_argument("-r", "--orientation", help="Page orientation. Valid values are portrait and landscape.", type=str, default='')
    parser.add_argument("-o", "--output", help="Output filename. Must be docx.", type=str, default='')

    # Print version
    parser.add_argument("-v", "--version", action="version", version='%(prog)s - Version 1.0')

    # Parse arguments
    args = parser.parse_args()
    return args


def try_file(inputfile):
    try:
        f = open(inputfile)
    except FileNotFoundError:
        print("File not found. Exiting.")
        sys.exit(1)


def get_pdf_size(inputfile, papersize):    
    if papersize == '': 
        input_pdf = PdfFileReader(open(inputfile, 'rb'))
        pdf_size = input_pdf.getPage(0).mediaBox
        if str(pdf_size) == 'RectangleObject([0, 0, 612, 1008])':
             papersize = "legal"             
        elif str(pdf_size) == 'RectangleObject([0, 0, 1008, 612])':
             papersize = "legal"             
        elif str(pdf_size) == 'RectangleObject([0, 0, 612, 792])':
             papersize = "letter"             
        elif str(pdf_size) == 'RectangleObject([0, 0, 792, 612])':
             papersize = "letter"             
        elif str(pdf_size) == 'RectangleObject([0, 0, 841.89, 595.276])':
             papersize = "a4"             
        elif str(pdf_size) == 'RectangleObject([0, 0, 595.276, 841.89])':
             papersize = "a4"             
        else: 
             papersize = "a4"             
        print("Papersize is", papersize)         
    else: 
        print("Papersize is", papersize)         
    global pdfSize
    pdfSize = papersize
    

def get_pdf_orientation(inputfile, orientation):
    if orientation == '': 
        pdf = PdfFileReader(inputfile)
        page = pdf.getPage(0).mediaBox
        if page.getUpperRight_x() - page.getUpperLeft_x() > page.getUpperRight_y() - page.getLowerRight_y():
            orientation = 'landscape'
        else:
            orientation = 'portrait'
        print("Orientation is", orientation)  
    else: 
        print("Orientation is", orientation)  
    global pdfOrientation
    pdfOrientation = orientation
    
                
def pdf_to_png(inputfile):
    global png_dir
    png_dir = "temp"
    try:
        os.mkdir(png_dir)
        print("Directory", png_dir, "created.") 
    except FileExistsError:
        print("Directory", png_dir, "already exists.")
    print("Converting pdf file to png.")
    global cur_dir
    cur_dir = os.getcwd()
    output_path = cur_dir + "/" + png_dir 
    global input_filename
    input_filename = os.path.splitext(os.path.basename(inputfile))[0]
    all_pages = Image(filename=inputfile, resolution=500)
    for i, page in enumerate(all_pages.sequence):
        with Image(page) as img:
            img.format = 'png'            
            image_filename = '{}-{}.png'.format(input_filename, i)
            print("Created", image_filename)
            image_filename = os.path.join(output_path, image_filename)
            img.save(filename=image_filename)
      
    
def png_to_docx(outputfile):
    papersize = pdfSize 
    orientation = pdfOrientation 
    templates_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = templates_dir + "/" + papersize + "-" + orientation + ".docx"
    width_dict = {'a4': 8.27, 'legal': 8.5, 'letter': 8.5}
    height_dict = {'a4': 11.69, 'legal': 14, 'letter': 11}    
    png_path = cur_dir + "/" + png_dir + '/*.png'
    if outputfile == '': 
        outputfile = input_filename + ".docx"  
    print("Inserting png files to", outputfile)
    document = Document(template_path)
    for pngfile in glob.glob(png_path):
        if orientation == 'portrait': 
        	document.add_picture(pngfile, width=Inches(width_dict[papersize]), height=Inches(height_dict[papersize]))
        else:
            document.add_picture(pngfile, width=Inches(height_dict[papersize]), height=Inches(width_dict[papersize]))        
    print("Saving", outputfile)
    document.save(outputfile)
    
    
def cleanup(): 
    print("Removing temp directory.")
    shutil.rmtree(png_dir, ignore_errors=True)
    print("Done!")
    return status
    


if __name__ == '__main__':
    args = parseArguments()
    try_file(args.inputPdf)
    get_pdf_size(args.inputPdf, args.papersize, )
    get_pdf_orientation(args.inputPdf, args.orientation)
    pdf_to_png(args.inputPdf) 
    png_to_docx(args.output)
    cleanup()
    
    
